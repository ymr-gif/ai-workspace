import asyncio
import logging
import re
import uuid
from pathlib import Path

from core.db import AsyncSessionLocal
from llm.embeddings import embed
from models import File, FileChunk
from observability.file_metrics import record_chunks

logger = logging.getLogger("processor")

CHUNK_SIZE    = 1600
CHUNK_OVERLAP = 200
_PROGRESS_TTL = 300  # Redis key TTL in seconds


def extract_text(storage_path: str, mime_type: str) -> str:
    path = Path(storage_path)
    mt   = (mime_type or "").lower()

    if mt == "application/pdf" or path.suffix.lower() == ".pdf":
        return _extract_pdf(path)

    if mt in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or path.suffix.lower() in (".docx", ".doc"):
        return _extract_docx(path)

    if mt in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ) or path.suffix.lower() in (".xlsx", ".xls"):
        return _extract_excel(path)

    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        logger.warning("[processor] text read failed path=%s err=%s", path, e)
        return ""


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except ImportError:
        logger.warning("[processor] pypdf not installed")
        return ""
    except Exception as e:
        logger.warning("[processor] pdf failed path=%s err=%s", path, e)
        return ""


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
        doc   = Document(str(path))
        parts = []

        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)

        for table in doc.tables:
            rows = []
            for i, row in enumerate(table.rows):
                seen: set[int] = set()
                cells = []
                for cell in row.cells:
                    tc_id = id(cell._tc)
                    if tc_id not in seen:
                        seen.add(tc_id)
                        cells.append(cell.text.replace("\n", " ").strip())
                if cells:
                    rows.append(" | ".join(cells))
                    if i == 0:
                        rows.append(" | ".join(["---"] * len(cells)))
            if rows:
                parts.append("\n".join(rows))

        return "\n\n".join(parts)
    except ImportError:
        logger.warning("[processor] python-docx not installed")
        return ""
    except Exception as e:
        logger.warning("[processor] docx failed path=%s err=%s", path, e)
        return ""


def _extract_excel(path: Path) -> str:
    try:
        import openpyxl
        wb    = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            rows = []
            for i, row in enumerate(sheet.iter_rows(values_only=True)):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    rows.append(" | ".join(cells))
                    if i == 0:
                        rows.append(" | ".join(["---"] * len(cells)))
            if rows:
                parts.append(f"## Sheet: {sheet.title}\n\n" + "\n".join(rows))
        wb.close()
        return "\n\n".join(parts)
    except ImportError:
        logger.warning("[processor] openpyxl not installed")
        return ""
    except Exception as e:
        logger.warning("[processor] excel failed path=%s err=%s", path, e)
        return ""


def _split_semantic(text: str) -> list[str]:
    """Paragraph → sentence → word split into atomic units ≤ CHUNK_SIZE."""
    units: list[str] = []
    for para in re.split(r"\n{2,}", text):
        para = para.strip()
        if not para:
            continue
        if len(para) <= CHUNK_SIZE:
            units.append(para)
            continue
        for sent in re.split(r"(?<=[.!?])\s+", para):
            sent = sent.strip()
            if not sent:
                continue
            if len(sent) <= CHUNK_SIZE:
                units.append(sent)
                continue
            buf: list[str] = []
            buf_len = 0
            for word in sent.split():
                if buf_len + len(word) + 1 > CHUNK_SIZE and buf:
                    units.append(" ".join(buf))
                    buf, buf_len = [], 0
                buf.append(word)
                buf_len += len(word) + 1
            if buf:
                units.append(" ".join(buf))
    return units


def _merge_with_overlap(units: list[str], size: int, overlap: int) -> list[str]:
    """Greedily merge atomic units into chunks with sentence-aligned overlap tail."""
    chunks: list[str] = []
    buf: list[str]    = []
    buf_len           = 0

    for unit in units:
        sep = 1 if buf else 0
        if buf and buf_len + sep + len(unit) > size:
            chunk = "\n".join(buf)
            chunks.append(chunk)
            tail = chunk[-overlap:] if len(chunk) > overlap else chunk
            m    = re.search(r"(?<=[.!?\n])\s*\S", tail)
            tail = tail[m.start():].strip() if m else tail.strip()
            buf     = [tail] if tail else []
            buf_len = len(tail)
        buf.append(unit)
        buf_len += sep + len(unit)

    if buf:
        chunk = "\n".join(buf).strip()
        if chunk:
            chunks.append(chunk)

    return chunks


def chunk_text(text: str) -> list[str]:
    if not text.strip():
        return []
    return _merge_with_overlap(_split_semantic(text), CHUNK_SIZE, CHUNK_OVERLAP)


async def extract_url_text(url: str) -> tuple[str, str]:
    """Fetch URL, extract readable text. Returns (text, title)."""
    import httpx
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        logger.warning("[processor] beautifulsoup4 not installed")
        return "", url

    try:
        async with httpx.AsyncClient(timeout=15, follow_redirects=True) as client:
            resp = await client.get(url, headers={"User-Agent": "Mozilla/5.0"})
            resp.raise_for_status()

        soup  = BeautifulSoup(resp.text, "lxml")
        title = (soup.find("title") or {}).get_text(strip=True) if soup.find("title") else url

        for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        return text, title

    except Exception as e:
        logger.warning("[processor] url fetch failed url=%s err=%s", url, e)
        return "", url


async def process_file_async(file_id: uuid.UUID, storage_path: str, mime_type: str) -> None:
    """Enqueue via ARQ when pool available; run inline as fallback (scheduler_worker)."""
    from core.arq_pool import get_arq_pool
    pool = get_arq_pool()
    if pool is not None:
        await pool.enqueue_job("process_file_job", str(file_id), storage_path, mime_type)
        return
    async with AsyncSessionLocal() as db:
        try:
            await _process(db, file_id, storage_path, mime_type)
        except Exception:
            logger.exception("[processor] failed file_id=%s", file_id)
            row = await db.get(File, file_id)
            if row:
                row.upload_status = "error"
                await db.commit()


async def _process(db, file_id: uuid.UUID, storage_path: str, mime_type: str) -> None:
    row = await db.get(File, file_id)
    if not row:
        return

    row.upload_status = "processing"
    await db.commit()

    text = extract_text(storage_path, mime_type)
    if not text.strip():
        row.upload_status = "error"
        await db.commit()
        logger.warning("[processor] empty text file_id=%s", file_id)
        return

    chunks = chunk_text(text)
    total  = len(chunks)
    saved  = 0
    pkey   = f"proc_progress:{file_id}"

    from core.redis_client import get_redis
    redis = None
    try:
        redis = get_redis()
    except Exception:
        pass

    # Embed all chunks concurrently (bounded by MAX_CONCURRENT_REQUESTS semaphore in llm/client.py)
    embeddings = await asyncio.gather(*[embed(c, input_type="passage") for c in chunks])

    for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):
        if emb:
            db.add(FileChunk(
                file_id     = file_id,
                chunk_index = i,
                content     = chunk,
                token_count = len(chunk.split()),
                embedding   = emb,
            ))
            saved += 1

        if redis and total > 0:
            try:
                await redis.setex(pkey, _PROGRESS_TTL, str(round((i + 1) / total, 3)))
            except Exception:
                redis = None

    row.upload_status = "ready"
    await db.commit()
    record_chunks(saved)
    logger.info("[processor] done file_id=%s chunks=%d/%d", file_id, saved, total)

    if redis:
        try:
            await redis.delete(pkey)
        except Exception:
            pass
